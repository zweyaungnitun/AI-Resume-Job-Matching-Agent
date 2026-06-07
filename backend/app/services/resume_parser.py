"""Resume parsing and extraction service"""

import json
import pdfplumber
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate

from app.config import settings


class ResumeParser:
    """Extract structured data from resume documents"""

    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(
            google_api_key=settings.GOOGLE_API_KEY,
            model=settings.GEMINI_MODEL,
            temperature=0.1
        )

    def parse_pdf(self, file_path: str) -> dict:
        """Parse PDF resume and extract text"""
        text = ""
        page_count = 0
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

        return {
            "raw_text": text.strip(),
            "pages": page_count,
            "file_type": "pdf"
        }

    def parse_docx(self, file_path: str) -> dict:
        """Parse DOCX resume and extract text including tables"""
        doc = Document(file_path)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"

        return {
            "raw_text": text.strip(),
            "pages": len(doc.paragraphs),
            "file_type": "docx"
        }

    def extract_skills(self, text: str) -> list[str]:
        """Extract skills from resume text using LLM"""
        prompt = PromptTemplate(
            input_variables=["resume_text"],
            template="""Extract all technical and professional skills from the following resume.
Return ONLY a JSON array of skill strings, nothing else.

Resume:
{resume_text}

Response (JSON array only):"""
        )

        chain = prompt | self.llm
        response = chain.invoke({"resume_text": text})

        try:
            return self._clean_json_response(response.content)
        except Exception:
            return []

    def extract_experience(self, text: str) -> list[dict]:
        """Extract work experience from resume text using LLM"""
        prompt = PromptTemplate(
            input_variables=["resume_text"],
            template="""Extract all work experience from the following resume.
For each position, extract: title, company, years (as number), and responsibilities (as list of strings).
Return ONLY a JSON array of objects with keys: title, company, years, responsibilities.

Resume:
{resume_text}

Response (JSON array only):"""
        )

        chain = prompt | self.llm
        response = chain.invoke({"resume_text": text})

        try:
            return self._clean_json_response(response.content)
        except Exception:
            return []

    def _clean_json_response(self, content: str):
        """Extract and parse JSON from LLM response"""
        if not content:
            return []
        
        content = content.strip()
        if content.startswith("```json"):
            content = content[7:-3]
        elif content.startswith("```"):
            content = content[3:-3]
        
        # Remove any lingering markdown or markers
        content = content.strip()
        
        try:
            data = json.loads(content)
            return data if isinstance(data, list) else []
        except json.JSONDecodeError:
            # Fallback: try to find the first [ and last ]
            start = content.find("[")
            end = content.rfind("]")
            if start != -1 and end != -1:
                try:
                    data = json.loads(content[start:end+1])
                    return data if isinstance(data, list) else []
                except json.JSONDecodeError:
                    pass
            return []

    def export_to_docx(self, resume_text: str, output_path: str) -> str:
        """Export resume text as DOCX file"""
        doc = Document()

        # Parse resume text into sections
        sections = resume_text.split('\n\n')

        for section in sections:
            if not section.strip():
                continue

            lines = section.strip().split('\n')
            for line in lines:
                if line.strip():
                    # Simple formatting: detect headers (all caps or short lines)
                    if len(line.strip()) < 50 and line.isupper():
                        heading = doc.add_heading(line.strip(), level=2)
                        heading.style = 'Heading 2'
                    else:
                        doc.add_paragraph(line.strip())

        doc.save(output_path)
        return output_path
